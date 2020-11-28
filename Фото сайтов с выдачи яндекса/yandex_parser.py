import os

from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from selenium.webdriver.chrome.options import Options


search = 'site: vsk.ru купить страховку'
file_name = 'Отчет по ключевой фразе'

files = os.listdir(os.getcwd())
if file_name + '.rtf' in files:
    os.remove(file_name + '.rtf')

document = Document()
chrome_options = Options()
chrome_options.add_argument("--headless")
driver = webdriver.Chrome(options=chrome_options)

driver.get(f'https://yandex.ru/search/?text={search}&lr=213')
source = BeautifulSoup(str(driver.page_source), 'html.parser')

title_name = source.find_all('div', attrs={'class': 'organic__url-text'})
description = source.find_all('div', attrs={'class': 'text-container typo typo_text_m typo_line_m organic__text'})
link = source.find_all('a', attrs={'class': 'link link_theme_outer path__item i-bem'})

try:
    for number in range(25):
        driver.get(link[number]['href'])
        driver.save_screenshot(f'my_screenshot{number}.png')
        # print(link[number]['href'])
        # print(title_name[number].text)
        # print(description[number].text)
        p = document.add_paragraph()
        runner = p.add_run(f'{number + 1}. ' + str(title_name[number].text) + '\n').bold = True

        document.add_paragraph(str(description[number].text) + '\n')
        document.add_picture(f'my_screenshot{number}.png', width=Inches(5.0))
        os.remove(f'my_screenshot{number}.png')
except IndexError:
    driver.quit()

document.save(file_name + '.docx')
base = os.path.splitext(file_name + '.docx')[0]
os.rename(file_name + '.docx', base + '.rtf')

